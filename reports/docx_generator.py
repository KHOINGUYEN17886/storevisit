import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from data.models import StoreReportData, FormPhoto

logger = logging.getLogger(__name__)

class DocxGenerator:
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def generate(self, data: StoreReportData) -> str:
        """
        Generates a professional Word document report for a store visit.
        """
        doc = Document()
        
        # Page setups: 1 inch margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Style setups
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Arial'
        style_normal.font.size = Pt(11)
        style_normal.font.color.rgb = RGBColor(44, 62, 80) # Slate text

        # Document Header / Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run("BIÊN BẢN ĐÁNH GIÁ CHẤT LƯỢNG CỬA HÀNG")
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(16)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(10, 35, 66) # Deep Slate Blue

        p_subtitle = doc.add_paragraph()
        p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_subtitle.add_run(f"Hệ thống StoreVisit Pro - Ngày báo cáo: {data.form_response.report_date if data.form_response else datetime.now().strftime('%d/%m/%Y')}")
        run_sub.font.italic = True
        run_sub.font.size = Pt(10)
        run_sub.font.color.rgb = RGBColor(127, 140, 141)

        # Metadata Table
        doc.add_heading("I. THÔNG TIN CHUNG", level=1)

        # Populate meta
        meta_data = [
            ("Mã cửa hàng", data.metadata.store_code),
            ("Tên cửa hàng", data.metadata.store_name),
            ("Khu vực / ASM", f"{data.metadata.region} / {data.metadata.asm_name}"),
            ("Cửa hàng trưởng / Ca trưởng", data.staff.cht_name),
            ("Ngày đánh giá", data.form_response.report_date if data.form_response else datetime.now().strftime('%Y-%m-%d'))
        ]

        # Đồng bộ 30-07: hiển thị loại kiểm tra (own/cross/opening) + thông tin khai trương nếu có
        inspection_mode = (data.form_response.inspection_mode if data.form_response else "own") or "own"
        MODE_LABELS = {"own": "Kiểm tra nội bộ", "cross": "Kiểm tra chéo", "opening": "Khai trương"}
        meta_data.append(("Loại kiểm tra", MODE_LABELS.get(inspection_mode, inspection_mode)))
        if inspection_mode == "opening" and data.form_response:
            OPENING_TYPE_LABELS = {"new": "Mở mới", "reopen": "Tái khai trương"}
            OPENING_PHASE_LABELS = {"before": "Trước khai trương", "day": "Ngày khai trương", "after": "Sau khai trương"}
            READINESS_LABELS = {"ready": "Sẵn sàng", "minor_fix": "Cần khắc phục nhỏ", "not_ready": "Chưa sẵn sàng"}
            meta_data.append(("Loại khai trương", OPENING_TYPE_LABELS.get(data.form_response.opening_type, data.form_response.opening_type or "-")))
            meta_data.append(("Giai đoạn", OPENING_PHASE_LABELS.get(data.form_response.opening_phase, data.form_response.opening_phase or "-")))
            meta_data.append(("Ngày khai trương chính thức", data.form_response.opening_date or "-"))
            meta_data.append(("Mức độ sẵn sàng", READINESS_LABELS.get(data.form_response.opening_readiness, data.form_response.opening_readiness or "-")))

        meta_table = doc.add_table(rows=len(meta_data), cols=2)
        meta_table.style = 'Table Grid'

        for i, (label, val) in enumerate(meta_data):
            # Bold label cell
            cell_lbl = meta_table.cell(i, 0)
            cell_lbl.text = label
            cell_lbl.paragraphs[0].runs[0].font.bold = True
            
            # Value cell
            meta_table.cell(i, 1).text = str(val)

        doc.add_paragraph() # spacer

        # Performance Summary Table
        doc.add_heading("II. HIỆU QUẢ HOẠT ĐỘNG (MTD)", level=1)
        rev = data.revenue
        perf_table = doc.add_table(rows=4, cols=2)
        perf_table.style = 'Table Grid'
        
        perf_data = [
            ("Doanh thu thực tế (MTD)", f"{rev.revenue_actual:,} VND"),
            ("Chỉ tiêu tháng", f"{rev.revenue_target:,} VND"),
            ("Tỷ lệ hoàn thành chỉ tiêu", f"{rev.attainment_pct:.1f}%"),
            ("So sánh với tháng trước (MoM)", f"{rev.mom_change_pct:+.1f}%")
        ]
        
        for i, (label, val) in enumerate(perf_data):
            cell_lbl = perf_table.cell(i, 0)
            cell_lbl.text = label
            cell_lbl.paragraphs[0].runs[0].font.bold = True
            
            perf_table.cell(i, 1).text = str(val)

        doc.add_paragraph() # spacer

        # Checklist Evaluations
        doc.add_heading("III. ĐÁNH GIÁ CÁC TIÊU CHÍ VẬN HÀNH", level=1)
        
        ratings_table = doc.add_table(rows=6, cols=2)
        ratings_table.style = 'Table Grid'
        
        # Header row
        hdr_cells = ratings_table.rows[0].cells
        hdr_cells[0].text = "Hạng mục đánh giá"
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].text = "Xếp loại"
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True
        
        # Paint header cell bg
        for cell in hdr_cells:
            shading = parse_xml(r'<w:shd {} w:fill="0A2342"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading)
            for r in cell.paragraphs[0].runs:
                r.font.color.rgb = RGBColor(255, 255, 255)

        categories = [
            ("1. Mặt tiền cửa hàng (Frontage)", data.frontage_rating or "Đạt"),
            ("2. Không gian bên trong (Inner space)", data.form_response.rating_inner if data.form_response else "Đạt"),
            ("3. Trưng bày hàng hóa (Merchandising)", data.form_response.rating_merch if data.form_response else "Đạt"),
            ("4. Nhân sự & Tác phong (Staff)", data.form_response.rating_staff if data.form_response else "Đạt"),
            ("5. Cơ sở vật chất & An ninh (CSVC)", data.form_response.rating_csvc if data.form_response else "Đạt")
        ]

        for i, (cat, rating) in enumerate(categories):
            row_idx = i + 1
            ratings_table.cell(row_idx, 0).text = cat
            cell_rating = ratings_table.cell(row_idx, 1)
            cell_rating.text = rating
            
            # Apply color based on rating
            run = cell_rating.paragraphs[0].runs[0]
            run.font.bold = True
            if rating.lower() in ["chưa đạt", "không đạt"]:
                run.font.color.rgb = RGBColor(192, 57, 43) # Red
            elif rating.lower() == "tốt":
                run.font.color.rgb = RGBColor(39, 174, 96) # Green
            else:
                run.font.color.rgb = RGBColor(243, 156, 18) # Amber

        doc.add_paragraph() # spacer

        # Detailed Issues & Action Plans
        doc.add_heading("IV. CHI TIẾT CÁC LỖI & PHƯƠNG ÁN KHẮC PHỤC", level=1)
        
        if not data.issues:
            p_no_issue = doc.add_paragraph()
            p_no_issue.add_run("Chúc mừng! Cửa hàng không có lỗi tồn đọng hoặc tất cả lỗi vận hành đều đạt tiêu chuẩn.").font.italic = True
        else:
            for issue in data.issues:
                p_issue_title = doc.add_paragraph()
                p_issue_title.paragraph_format.space_before = Pt(12)
                p_issue_title.paragraph_format.keep_with_next = True
                
                run_idx = p_issue_title.add_run(f"Lỗi {issue.index}: {issue.label}")
                run_idx.font.bold = True
                run_idx.font.size = Pt(12)
                run_idx.font.color.rgb = RGBColor(192, 57, 43)
                
                p_detail = doc.add_paragraph()
                p_detail.paragraph_format.left_indent = Inches(0.25)
                p_detail.add_run("• Hiện trạng lỗi: ").font.bold = True
                p_detail.add_run(f"{issue.issue}\n")
                p_detail.add_run("• Người chịu trách nhiệm: ").font.bold = True
                p_detail.add_run(f"{issue.assignee}\n")
                p_detail.add_run("• Trạng thái khắc phục: ").font.bold = True
                p_detail.add_run(f"{issue.status}\n")
                p_detail.add_run("• Chi tiết kế hoạch & Hạn xử lý: ").font.bold = True
                p_detail.add_run(f"{issue.notes}\n")

                # Try to add before / after photo if available
                # Search for photos in form_response
                if data.form_response:
                    label_to_sec_key = {
                        "Mặt tiền": "frontage",
                        "Không gian trong": "inner",
                        "Trưng bày AP": "merch_ap",
                        "Trưng bày PIE": "merch_pie",
                        "Trưng bày AB": "merch_ab",
                        "Trưng bày Anamai": "merch_anamai",
                        "Trưng bày Bonjour": "merch_bonjour",
                        "Phụ kiện": "merch_pk",
                        "Kho/Phòng thử": "warehouse",
                        "Kho hàng": "stockroom",
                        "Phòng thử đồ": "fitting_room",
                        "Nhà vệ sinh": "toilet",
                        "PCCC & Thoát hiểm": "fire_safety",
                        "Thu ngân": "cashier",
                        "Bao bì & An ninh": "packaging_security",
                        "Nhân sự & Tác phong": "staff",
                        "Bảo vệ": "security_guard"
                    }
                    sec_key = label_to_sec_key.get(issue.label)
                    if sec_key:
                        before_photos = [p for p in data.form_response.photos if p.section == f"issue_{sec_key}_before" and p.local_path and os.path.exists(p.local_path)]
                        after_photos = [p for p in data.form_response.photos if p.section == f"issue_{sec_key}_after" and p.local_path and os.path.exists(p.local_path)]
                        
                        if before_photos or after_photos:
                            # Add spacing
                            doc.add_paragraph()
                            
                            if before_photos and after_photos:
                                # Two columns table for before / after
                                table = doc.add_table(rows=1, cols=2)
                                table.autofit = False
                                
                                # Before column
                                cell_before = table.cell(0, 0)
                                cell_before.width = Inches(3.0)
                                p_bef = cell_before.paragraphs[0]
                                p_bef.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_bef.add_run("Hình ảnh lỗi ghi nhận:\n").font.bold = True
                                try:
                                    p_bef.add_run().add_picture(before_photos[0].local_path, width=Inches(2.8))
                                except Exception as img_err:
                                    p_bef.add_run(f"[Không thể tải ảnh: {img_err}]").font.italic = True
                                    
                                # After column
                                cell_after = table.cell(0, 1)
                                cell_after.width = Inches(3.0)
                                p_aft = cell_after.paragraphs[0]
                                p_aft.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_aft.add_run("Hình ảnh sau khắc phục:\n").font.bold = True
                                try:
                                    p_aft.add_run().add_picture(after_photos[0].local_path, width=Inches(2.8))
                                except Exception as img_err:
                                    p_aft.add_run(f"[Không thể tải ảnh: {img_err}]").font.italic = True
                            elif before_photos:
                                # Only before photo
                                p_img = doc.add_paragraph()
                                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_img.add_run("Hình ảnh lỗi ghi nhận:\n").font.bold = True
                                try:
                                    p_img.add_run().add_picture(before_photos[0].local_path, width=Inches(3.5))
                                except Exception as img_err:
                                    p_img.add_run(f"[Không thể tải ảnh: {img_err}]").font.italic = True
                            elif after_photos:
                                # Only after photo
                                p_img = doc.add_paragraph()
                                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_img.add_run("Hình ảnh sau khắc phục:\n").font.bold = True
                                try:
                                    p_img.add_run().add_picture(after_photos[0].local_path, width=Inches(3.5))
                                except Exception as img_err:
                                    p_img.add_run(f"[Không thể tải ảnh: {img_err}]").font.italic = True
                            doc.add_paragraph() # spacer

        # Save doc
        from utils.filename_formatter import format_store_output_filename
        r_date = data.form_response.report_date if data.form_response else ""
        filename = format_store_output_filename(data.metadata.store_name, data.metadata.asm_name, r_date, "docx")
        dest_path = os.path.join(self.output_dir, filename)
        doc.save(dest_path)
        logger.info(f"Generated docx report successfully at: {dest_path}")
        return dest_path
